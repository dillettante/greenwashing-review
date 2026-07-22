from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .analysis import PATTERN_LABELS

BODY_FONT = "바탕"
HEADING_FONT = "돋움"


def _font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr().rFonts
    rpr.set(qn("w:eastAsia"), name)
    rpr.set(qn("w:ascii"), name)
    rpr.set(qn("w:hAnsi"), name)


def _kv(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def create_assessment_report_docx(result: dict[str, Any], authorities: dict[str, dict[str, Any]], output_path: Path) -> None:
    """3-legal-review-report.md와 동일 내용을 house-style .docx로 생성한다."""
    ctx = result["context"]
    claims = result["claims"]
    doc = Document()
    section = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Mm(30))
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = 1.6
    for name in ("Heading 1", "Heading 2"):
        st = doc.styles[name]
        st.font.name = HEADING_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        st.font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(title.add_run("그린워싱 법률검토보고서"), HEADING_FONT, 16, True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(sub.add_run(f"사건: {result['matter_id']} | 작성: {result['created_at'][:10]}"), BODY_FONT, 11)
    doc.add_paragraph("변호사 검토용 초안. 위험점수는 우선순위 도구이며 위법성의 자동 결론이 아닙니다.")

    es = result.get("exec_summary") or {}
    if es:
        doc.add_heading("경영진 요약 (Executive Summary)", level=1)
        if es.get("headline"):
            para = doc.add_paragraph()
            _font(para.add_run(es["headline"]), BODY_FONT, 11, True)
        _bullets(doc, es.get("findings") or [])
        if es.get("worst_case"):
            doc.add_paragraph(f"최대 리스크 시나리오: {es['worst_case']}")
        if es.get("recommendation"):
            doc.add_paragraph(f"권고: {es['recommendation']}")

    doc.add_heading("1. 검토 대상 및 전제", level=1)
    _kv(doc, [
        ("기업", ctx.get("company", "[확인 필요]")),
        ("제품·서비스", ctx.get("product", "[확인 필요]")),
        ("매체", ctx.get("medium", "[확인 필요]")),
        ("예상 독자", ctx.get("audience", "[확인 필요]")),
        ("게시일", ctx.get("published_date", "[확인 필요]")),
    ])
    if result.get("warnings"):
        doc.add_heading("확인 필요 사항", level=2)
        _bullets(doc, result["warnings"])

    evaluated = [c for c in claims if c.get("evaluation")]
    narratives = result.get("narratives") or []
    gateway = result.get("gateway") or {}
    section = 2

    doc.add_heading(f"{section}. 결론 요약", level=1)
    if narratives:
        axes = " / ".join(f"[{i}] {n.get('axis','')}" for i, n in enumerate(narratives, 1))
        doc.add_paragraph(f"핵심 위험 축 {len(narratives)}개 — {axes}. 상세는 '종합 위험 분석'.")
    if evaluated:
        af = {"있음": 0, "불확실": 0, "없음": 0}
        rf = {"매우 높음": 0, "높음": 0, "중간": 0, "낮음": 0}
        for c in evaluated:
            ev = c["evaluation"]
            af[ev.get("applicability_final", "불확실")] = af.get(ev.get("applicability_final", "불확실"), 0) + 1
            if ev.get("risk_final") in rf:
                rf[ev["risk_final"]] += 1
        doc.add_paragraph(
            f"[정밀평가 최종 분포] {len(evaluated)}건 — 광고성 있음 {af['있음']} · 불확실 {af['불확실']} · 없음 {af['없음']} / "
            f"위험 매우 높음 {rf['매우 높음']} · 높음 {rf['높음']} · 중간 {rf['중간']} · 낮음 {rf['낮음']}."
        )
    else:
        counts = {"매우 높음": 0, "높음": 0, "중간": 0, "낮음": 0}
        for c in claims:
            counts[c["risk_band"]] += 1
        doc.add_paragraph(
            f"[기계 트리아지 분포 — 참고용] 환경 주장 {len(claims)}건 — 매우 높음 {counts['매우 높음']} · 높음 {counts['높음']} · "
            f"중간 {counts['중간']} · 낮음 {counts['낮음']}. 정규식 값(법적 판단 아님)."
        )

    if narratives:
        section += 1
        doc.add_heading(f"{section}. 종합 위험 분석 — 사건 서사", level=1)
        for i, n in enumerate(narratives, 1):
            doc.add_heading(f"축 {i}. {n.get('axis','')}", level=2)
            rows = [(label, n[key]) for label, key in
                    [("회사의 서사", "company_story"), ("확인된 사실", "confirmed_reality"),
                     ("괴리·법적 의미", "gap"), ("법적 평가", "legal_significance")] if n.get(key)]
            if rows:
                _kv(doc, rows)
            if n.get("claim_ids"):
                doc.add_paragraph("소속 주장: " + ", ".join(n["claim_ids"]))

    if gateway:
        section += 1
        doc.add_heading(f"{section}. 관문 쟁점 — 광고 해당성·경로 선택", level=1)
        ad = gateway.get("ad_applicability") or {}
        if ad:
            doc.add_heading("광고 해당성 (전 주장 공통 선결문제)", level=2)
            if ad.get("analysis"):
                doc.add_paragraph(ad["analysis"])
            _bullets(doc, [f"{m.get('medium','')}: {m.get('conclusion','')} — {m.get('reason','')}"
                           for m in ad.get("media") or []])
            _bullets(doc, [f"근거: {pr.get('cite','')}" + (f" [{pr['status']}]" if pr.get("status") else "")
                           + (f" — {pr['holding']}" if pr.get("holding") else "")
                           for pr in ad.get("precedents") or []])
            if ad.get("conclusion"):
                doc.add_paragraph(f"결론: {ad['conclusion']}")
        routes_alt = gateway.get("alternative_routes") or []
        if routes_alt:
            doc.add_heading("대안 경로 비교", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for j, head in enumerate(("경로", "요건", "제재·효과", "실익·한계")):
                table.rows[0].cells[j].text = head
            for r in routes_alt:
                cells = table.add_row().cells
                for j, key in enumerate(("route", "requirements", "sanctions", "pros_cons")):
                    cells[j].text = str(r.get(key, ""))

    exposure = result.get("exposure") or {}
    if exposure:
        section += 1
        doc.add_heading(f"{section}. 정량 리스크·제재 전망", level=1)
        sanctions = exposure.get("sanctions") or []
        if sanctions:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for j, head in enumerate(("경로", "근거", "노출(상한·구조)", "유사 사건 벤치마크")):
                table.rows[0].cells[j].text = head
            for s in sanctions:
                cells = table.add_row().cells
                for j, key in enumerate(("route", "basis", "exposure", "benchmark")):
                    cells[j].text = str(s.get(key, ""))
        if exposure.get("derivative_risks"):
            doc.add_paragraph("파생 리스크:")
            _bullets(doc, exposure["derivative_risks"])
        if exposure.get("caveat"):
            doc.add_paragraph(exposure["caveat"])

    detailed = evaluated if evaluated else [c for c in claims if c["applicability"] == "있음"][:12]
    section += 1
    claim_sec = section
    doc.add_heading(f"{claim_sec}. 주장별 검토", level=1)
    for idx, c in enumerate(detailed, 1):
        ev = c.get("evaluation") or {}
        final_risk = ev.get("risk_final") or c["risk_band"]
        doc.add_heading(f"{claim_sec}-{idx}. {c['claim_id']} — 최종 위험 {final_risk}", level=2)
        doc.add_paragraph(f"원문: “{c['quote']}”")
        anchor = c.get("anchor") or {}
        rows = [("위치", f"{c['filename']} {c['page']}쪽")]
        if anchor.get("status") == "not_found":
            rows.append(("⚠️ 원문 앵커", "인용 원문 미확인 — 재확인 전 인용 금지"))
        if ev.get("applicability_final"):
            rows.append(("광고성(최종)", ev["applicability_final"]))
        rows += [("주장 대상", c["subject_scope"]),
                 ("유형", ", ".join(PATTERN_LABELS.get(p, p) for p in c["patterns"]))]
        if c.get("why_flagged"):
            rows.append(("선별 사유(세션 추출)", c["why_flagged"]))
        if not ev:
            rows.append(("[미평가] 기계 1차분류(법적 판단 아님)", c["legal_call"]))
        _kv(doc, rows)
        if ev:
            provs = ev.get("provisions") or []
            if provs:
                doc.add_paragraph("법적 평가 — 적용 조문(포섭):")
                _bullets(doc, [f"{p.get('authority_id','')} {p.get('cite','')}" + (f" — {p['label']}" if p.get("label") else "") for p in provs])
            if ev.get("assessment"):
                doc.add_paragraph(f"포섭·판단: {ev['assessment']}")
            if ev.get("misleading"):
                doc.add_paragraph(f"오인가능성: {ev['misleading']}")
            precs = ev.get("precedents") or []
            doc.add_paragraph("참조 심결례·판례: " + ("; ".join(f"{pr.get('cite','')}" + (f"[{pr['status']}]" if pr.get("status") else "") for pr in precs) if precs else "로컬 의결서 아카이브(corpus/raw/KR/cases/_index.csv)에서 유사 표시·광고 의결서 검색·대조"))
            ver = ev.get("verification")
            if ver:
                doc.add_paragraph(f"실증·검증(웹): [{ver.get('verdict','미확인')}] {ver.get('summary','')}")
                for src in ver.get("sources", []):
                    doc.add_paragraph(f"[{src.get('stance','중립')}] {src.get('title','')} ({src.get('publisher','')} {src.get('date','')}) {src.get('url','')} — {src.get('finding','')}", style="List Bullet")
            else:
                doc.add_paragraph("실증·검증(웹): [미실시] 회사 주장 지표의 실재·범위·반증 확인 필요")
            if ev.get("confirm_needed"):
                doc.add_paragraph("[확인 필요]:")
                _bullets(doc, ev["confirm_needed"])
        if c["missing_evidence"] and not ev:
            doc.add_paragraph("추가 확보 필요 자료:")
            _bullets(doc, c["missing_evidence"])

    # 제출 경로 섹션은 두지 않는다(md와 동일) — 관문 쟁점의 '대안 경로 비교'와
    # '정량 리스크·제재 전망'이 법적 근거와 함께 다루므로 중복이다.
    section += 1
    doc.add_heading(f"{section}. 최종 검증 게이트", level=1)
    _bullets(doc, [
        "사건 당시 시행 법령과 현행 법령을 구분하여 재확인",
        "판례·심결례의 원문·절차단계 및 확정 여부 확인",
        "각 주장과 증거의 페이지·파일 해시 대조",
        "회사 환경성과·지표의 실재 여부 웹·원자료 재검증",
        "행위자·게시기간·도달범위·시정 여부 확인",
        "표시·광고 해당성과 소비자 오인 가능성(경로 공통 선결)",
        "환경기술 및 환경산업 지원법상 제품·행위자 범위",
        "실증자료 요청·보전 필요성",
        "행정조사와 형사절차의 순서 및 중복 위험",
        "고발인·피고발인 인적사항, 관할, 시효",
        "변호사 최종 승인 후 제출문서 확정",
    ])

    cited: dict[tuple[str, str], dict[str, Any]] = {}
    for c in detailed:
        for cit in c.get("legal_citations", []):
            cited[(cit["authority_id"], cit["provision_no"])] = cit
    if cited:
        section += 1
        doc.add_heading(f"{section}. 직접 근거 원문·버전", level=1)
        for cit in cited.values():
            doc.add_heading(f"{cit['title']} {cit['provision_no']} {cit.get('heading') or ''}", level=2)
            excerpt = cit["text"][:1200]
            if len(cit["text"]) > len(excerpt):
                excerpt += " […이하 로컬 조문 DB]"
            doc.add_paragraph(excerpt)
            doc.add_paragraph(f"시행일 {cit.get('effective_date') or '[확인 필요]'} · 조문 SHA-256 {cit['provision_sha256']}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
